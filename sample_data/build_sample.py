"""
Builds a synthetic "ticket log" .docx (mirrors what a real support ticket
export looks like: a table plus some free-text notes) with PII values we
choose ourselves, and writes out ground_truth.json listing every PII
instance we deliberately planted.

Because ground truth is generated from the exact same values used to build
the document, this gives an authoritative (not manually-annotated, so
error-free) test set for validating the redactor and computing
precision/recall/accuracy in evaluate.py.

Run:
    python build_sample.py
"""

import json
import os

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))

# Each row: (ticket_id, name, email, phone, company, issue_text, address, dob, notes)
ROWS = [
    dict(
        ticket_id="TCK-1001",
        name="Rashi Patil",
        email="rashi.patil@gmail.com",
        phone="+91 9876543210",
        company="Bluewave Technologies Pvt Ltd",
        issue="Order #4482 not delivered to my address at 221 MG Road, Koramangala, Bengaluru 560034.",
        dob_line="DOB: 14/03/1994",
        notes="Called from IP 103.25.14.90. Card charged was 4539 1488 0343 6467 for the refund.",
    ),
    dict(
        ticket_id="TCK-1002",
        name="Rohan Dey",
        email="rohan.dey@gmail.com",
        phone="+91 9123456789",
        company="Orion Data Systems Inc",
        issue="Unable to log in to my account since yesterday.",
        dob_line="Date of Birth: 02 July 1989",
        notes="SSN on file for verification: 219-09-9999. Support agent Dr. Karan Mehta assisted.",
    ),
    dict(
        ticket_id="TCK-1003",
        name="Ananya Sharma",
        email="ananya.sharma88@yahoo.com",
        phone="9988776655",
        company="Vertex Cloud Solutions",
        issue="Billing address needs to change to 45 Park Street, Salt Lake, Kolkata 700091.",
        dob_line="Born on 1990-11-23",
        notes="Reported by Mr. Vikram Rao on behalf of the customer. Connection logged from 192.168.10.55.",
    ),
]


def build_docx(path: str):
    doc = Document()
    doc.add_heading("Customer Support Ticket Log", level=1)
    doc.add_paragraph(
        "Weekly export of open tickets for the billing and account-access queue."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Ticket ID"
    hdr[1].text = "Customer Name"
    hdr[2].text = "Email"
    hdr[3].text = "Phone"
    hdr[4].text = "Company"

    for row in ROWS:
        cells = table.add_row().cells
        cells[0].text = row["ticket_id"]
        cells[1].text = row["name"]
        cells[2].text = row["email"]
        cells[3].text = row["phone"]
        cells[4].text = row["company"]

    doc.add_heading("Ticket Details", level=2)
    for row in ROWS:
        doc.add_paragraph(f"{row['ticket_id']} - Customer Name: {row['name']}", style="List Bullet")
        doc.add_paragraph(row["issue"])
        doc.add_paragraph(row["dob_line"])
        doc.add_paragraph(row["notes"])
        doc.add_paragraph("")

    doc.save(path)


def build_ground_truth(path: str):
    instances = []
    for row in ROWS:
        instances.append({"type": "NAME", "text": row["name"]})
        # name also appears a second time in the "Customer Name:" line
        instances.append({"type": "NAME", "text": row["name"]})
        instances.append({"type": "EMAIL", "text": row["email"]})
        instances.append({"type": "PHONE", "text": row["phone"]})
        instances.append({"type": "COMPANY", "text": row["company"]})

    # per-row planted extras (address / dob / ssn / cc / ip / extra name)
    instances += [
        {"type": "ADDRESS", "text": "221 MG Road, Koramangala, Bengaluru 560034"},
        {"type": "DOB", "text": "14/03/1994"},
        {"type": "IP", "text": "103.25.14.90"},
        {"type": "CREDIT_CARD", "text": "4539 1488 0343 6467"},

        {"type": "DOB", "text": "02 July 1989"},
        {"type": "SSN", "text": "219-09-9999"},
        {"type": "NAME", "text": "Karan Mehta"},

        {"type": "ADDRESS", "text": "45 Park Street, Salt Lake, Kolkata 700091"},
        {"type": "DOB", "text": "1990-11-23"},
        {"type": "NAME", "text": "Vikram Rao"},
        {"type": "IP", "text": "192.168.10.55"},
    ]

    with open(path, "w", encoding="utf-8") as f:
        json.dump(instances, f, indent=2)


if __name__ == "__main__":
    docx_path = os.path.join(HERE, "ticket_log_sample.docx")
    gt_path = os.path.join(HERE, "ground_truth.json")
    build_docx(docx_path)
    build_ground_truth(gt_path)
    print(f"Wrote {docx_path}")
    print(f"Wrote {gt_path}")
