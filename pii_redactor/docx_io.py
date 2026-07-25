"""
.docx reading/writing.

Word documents split a paragraph's visible text across many `<w:r>` runs
(spell-check boundaries, revision markers, etc.), so PII detection has to
happen on the paragraph's FULL joined text, not run-by-run. This module
joins each paragraph to plain text for detection, then writes the redacted
text back into the paragraph's first run (clearing the rest).

Trade-off (documented in README): this collapses a paragraph to a single
run, so run-level formatting that varied *within* a paragraph (e.g. one
word bolded) is not perfectly preserved -- only the first run's formatting
survives. Paragraph-level styling (heading style, alignment, table
structure) is fully preserved.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Iterator, List

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .detectors import Span
from .redactor import Redactor


def _iter_table_paragraphs(table: Table, _seen_tc=None) -> Iterator[Paragraph]:
    # python-docx's `row.cells` returns one entry PER GRID COLUMN a cell
    # occupies -- a cell horizontally merged across 3 columns is yielded
    # 3 times (same underlying <w:tc> element each time). Undeduped, this
    # redacts the same text repeatedly and, worse, feeds each pass's fake
    # replacement back in as if it were fresh real text to detect PII in
    # -- which is how a real Faker-generated placeholder ("Pearson PLC")
    # ended up looking like a "detected" company in early testing on the
    # real reference document, despite that string not existing anywhere
    # in the source file.
    #
    # Dedup by keeping the actual `<w:tc>` ELEMENT OBJECTS in a set/list
    # (not `id(cell._tc)` integers): python-docx/lxml can create a fresh,
    # short-lived proxy object each time a cell is accessed, and once the
    # previous proxy is garbage-collected Python is free to reuse its
    # memory address for something else entirely -- so bare `id()` values
    # can *collide between genuinely different cells*. This was caught by
    # a regression test suddenly losing paragraphs after an id()-based
    # version of this dedup was first tried. Holding a real reference in
    # `_seen_tc` keeps each object alive for the duration of the scan,
    # which makes identity comparison safe.
    if _seen_tc is None:
        _seen_tc = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in _seen_tc:
                continue
            _seen_tc.add(cell._tc)  # keeps the element alive -> id() stays unique
            yield from _iter_cell_paragraphs(cell, _seen_tc)


def _iter_cell_paragraphs(cell: _Cell, _seen_tc=None) -> Iterator[Paragraph]:
    for p in cell.paragraphs:
        yield p
    for nested in cell.tables:
        yield from _iter_table_paragraphs(nested, _seen_tc)


def _iter_all_paragraphs_raw(document: Document) -> Iterator[Paragraph]:
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        yield from _iter_table_paragraphs(table)

    # Same "keep a live reference, don't compare bare id()s" reasoning as
    # _iter_table_paragraphs above -- see that function's docstring.
    seen_parts = set()
    for section in document.sections:
        for part in (section.header, section.footer):
            if part.part in seen_parts:
                continue
            seen_parts.add(part.part)
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def iter_all_paragraphs(document: Document) -> Iterator[Paragraph]:
    """All paragraphs in the document body, inside tables (recursively),
    and in every section's headers/footers -- each physical paragraph
    exactly once.

    IMPORTANT, found while validating against the real reference document:
    two independent sources of duplication exist in a real .docx and both
    matter a lot for a redaction tool specifically, because reprocessing
    an already-redacted paragraph doesn't just waste time -- it detects
    PII inside the PREVIOUS pass's fake replacement text and redacts that
    "again", compounding into fake data that resembles nothing in the
    source (this surfaced as plausible-looking companies like "Pearson
    PLC" in the detected-span log that, on inspection, appeared nowhere
    in the source document at all):
      1. Multiple sections commonly share one underlying header/footer
         part ("Link to Previous" in Word); naively looping
         `document.sections` yields that shared part once per section.
      2. `row.cells` returns one entry per GRID COLUMN a cell spans, so
         a horizontally-merged cell (common in cover-page/contact-info
         tables) is yielded once per column it covers.
    `_iter_table_paragraphs` already dedups (2) within a single table
    pass; this function adds a final, unconditional dedup on the
    underlying XML element's identity as a backstop, since the two
    sources can still interact (a merged cell inside a shared header).
    """
    seen = set()
    for p in _iter_all_paragraphs_raw(document):
        if p._p in seen:
            continue
        seen.add(p._p)
        yield p


def _set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if not paragraph.runs:
        if new_text:
            paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


@dataclass
class DocxRedactionReport:
    paragraph_count: int = 0
    redacted_paragraph_count: int = 0
    spans_by_type: dict = field(default_factory=dict)
    all_spans: List[Span] = field(default_factory=list)

    def record(self, spans: List[Span]) -> None:
        for s in spans:
            self.spans_by_type[s.type] = self.spans_by_type.get(s.type, 0) + 1
        self.all_spans.extend(spans)


def redact_docx(input_path: str, output_path: str, seed: int = 42) -> DocxRedactionReport:
    document = Document(input_path)
    redactor = Redactor(seed=seed)
    report = DocxRedactionReport()

    for paragraph in iter_all_paragraphs(document):
        text = paragraph.text
        if not text.strip():
            continue
        report.paragraph_count += 1
        result = redactor.redact(text)
        if result.spans:
            report.redacted_paragraph_count += 1
            report.record(result.spans)
            _set_paragraph_text(paragraph, result.redacted_text)

    document.save(output_path)
    return report


def extract_text(input_path: str) -> str:
    """Plain-text dump of a .docx, paragraph-per-line (used for evaluation
    scripts / quick inspection, not for the redaction pipeline itself)."""
    document = Document(input_path)
    lines = []
    for paragraph in iter_all_paragraphs(document):
        if paragraph.text.strip():
            lines.append(paragraph.text)
    return "\n".join(lines)
